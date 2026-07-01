#!/usr/bin/env python3
"""
Gerador de documentos Word a partir de Template Mestre e ProcessData.

Este módulo preenche placeholders em DOCX templates com dados extratos.
Suporta dois formatos de placeholders:
- Colchetes: [PROC], [AI], [DOC]
- Chevrons: <<PROC>>, <<AI>>, <<DOC>>
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from src.models import ProcessData


class WordGeneratorError(Exception):
    """Erro durante geração de documento Word."""

    pass


class WordGenerator:
    """
    Gerador de DOCX a partir de Template Mestre.
    
    Substitui placeholders nos parágrafos e tabelas com dados do ProcessData.
    Preserva formatação original (negrito, itálico, fonte, etc).
    
    Suporta dois formatos de placeholders:
    - Colchetes: [PROC], [AI], [DOC]
    - Chevrons: <<PROC>>, <<AI>>, <<DOC>>
    """

    # Mapa de placeholders para campos de ProcessData
    PLACEHOLDER_MAP = {
        # Identificação do Processo
        "PROC": "process_number",
        "AI": "infraction_number",
        
        # Identificação do Contribuinte
        "NOME": "nome",
        "DOC": "cpf_cnpj",
        
        # Endereço do Contribuinte
        "LOG": "logradouro",
        "NR": "numero",
        "COMP": "complemento",
        "BAIRRO": "bairro",
        "MUN": "municipio",
        "UF": "uf",
        "CEP": "cep",
        
        # Contato do Contribuinte
        "FONE": "telefone",
        "EMAIL": "email",
    }

    def generate(
        self,
        template_path: Path,
        process_data: ProcessData,
        output_path: Path,
    ) -> None:
        """
        Gera DOCX preenchido a partir do template.
        
        Substitui placeholders ([PROC], <<PROC>>, etc) com dados.
        Preserva formatação original do template.
        
        Args:
            template_path: Caminho do template DOCX
            process_data: Dados para preenchimento
            output_path: Caminho para salvar DOCX gerado
            
        Raises:
            WordGeneratorError: Se template não existir ou geração falhar
        """
        if not template_path.exists():
            raise WordGeneratorError(f"Template não encontrado: {template_path}")

        try:
            # Abrir documento
            doc = Document(str(template_path))

            # Processar parágrafos
            for paragraph in doc.paragraphs:
                self._replace_placeholders_in_paragraph(paragraph, process_data)

            # Processar tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_in_paragraph(
                                paragraph, process_data
                            )

            # Salvar documento
            doc.save(str(output_path))

        except WordGeneratorError:
            raise
        except Exception as e:
            raise WordGeneratorError(f"Erro ao gerar documento: {e}") from e

    def _replace_placeholders_in_paragraph(
        self, paragraph, process_data: ProcessData
    ) -> None:
        """
        Substitui placeholders em um parágrafo.
        
        Itera sobre runs do parágrafo e substitui placeholders diretamente.
        Preserva formatação (negrito, itálico, fonte, etc).
        
        Suporta dois formatos:
        - Colchetes: [PROC], [AI], [DOC]
        - Chevrons: <<PROC>>, <<AI>>, <<DOC>>
        
        Args:
            paragraph: Parágrafo a processar
            process_data: Dados para substituição
        """
        for run in paragraph.runs:
            # Substituir cada placeholder no run
            for placeholder, field_name in self.PLACEHOLDER_MAP.items():
                value = getattr(process_data, field_name, "")
                
                # Formato com colchetes: [PROC]
                bracket_placeholder = f"[{placeholder}]"
                if bracket_placeholder in run.text:
                    run.text = run.text.replace(bracket_placeholder, value)
                
                # Formato com chevrons: <<PROC>>
                chevron_placeholder = f"<<{placeholder}>>"
                if chevron_placeholder in run.text:
                    run.text = run.text.replace(chevron_placeholder, value)
